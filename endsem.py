import streamlit as st
import re, random
from io import BytesIO
from copy import deepcopy
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import zipfile

# -----------------------
# Regex & XML Helpers
# -----------------------
CO_DESC_RE = re.compile(r'CO(\d)[:\s]+(.*)', re.I)
NUM_PREFIX_RE = re.compile(r'^\s*(\d+)\s*[\.\)]')
BLOOM_RE = re.compile(r'K\s*([1-6])', re.I)

def clean_xml_for_images(element):
    """Removes drawing and picture tags using local-name to avoid namespace errors."""
    for drawing in element.xpath(".//*[local-name()='drawing']"):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)
    for pic in element.xpath(".//*[local-name()='pic']"):
        parent = pic.getparent()
        if parent is not None:
            parent.remove(pic)

def replace_cell_with_cell(target_cell, src_cell):
    """Copies text/formatting but strips original image tags."""
    t_tc = target_cell._tc
    s_tc = deepcopy(src_cell._tc)
    
    clean_xml_for_images(s_tc)
    
    original_props = t_tc.get_or_add_tcPr()
    v_align = original_props.find(qn('w:vAlign'))
    
    for child in list(t_tc):
        if child.tag != qn('w:tcPr'):
            t_tc.remove(child)
            
    for child in list(s_tc):
        if child.tag != qn('w:tcPr'):
            t_tc.append(child)
            
    if v_align is not None:
        new_props = t_tc.get_or_add_tcPr()
        new_props.append(deepcopy(v_align))

def extract_images_from_cell(cell):
    images = []
    for blip in cell._element.xpath(".//*[local-name()='blip']"):
        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rId:
            try:
                part = cell.part.related_parts[rId]
                images.append(BytesIO(part.blob))
            except:
                continue
    return images

# -----------------------
# Core Logic
# -----------------------
def extract_bank_data(uploaded_file):
    doc = Document(BytesIO(uploaded_file.read()))
    questions = []
    co_map = {}
    
    for table in doc.tables:
        rows = table.rows
        for i, row in enumerate(rows):
            cells_text = [c.text.strip() for c in row.cells]
            text_blob = " ".join(cells_text)
            
            co_matches = CO_DESC_RE.findall(text_blob)
            for num, desc in co_matches:
                if num not in co_map: co_map[num] = desc.strip()
            
            tag = None
            for txt in cells_text:
                if re.match(r'^[1-5][ABC]$', txt.upper()):
                    tag = txt.upper()
                    break
            
            if tag and "ans" not in cells_text[0].lower():
                bm = BLOOM_RE.search(text_blob)
                main_idx = max(range(len(row.cells)), key=lambda idx: len(row.cells[idx].text))
                imgs = extract_images_from_cell(row.cells[main_idx])
                
                ans_data = None
                if i + 1 < len(rows):
                    next_row_cells = [c.text.strip() for c in rows[i+1].cells]
                    if "ans" in next_row_cells[0].lower():
                        ans_idx = max(range(len(rows[i+1].cells)), key=lambda idx: len(rows[i+1].cells[idx].text))
                        ans_data = rows[i+1].cells[ans_idx].text

                questions.append({
                    "id": random.random(),
                    "tag": tag,
                    "bloom": f"K{bm.group(1)}" if bm else "",
                    "q_cell": row.cells[main_idx], 
                    "ans_text": ans_data,
                    "images": imgs
                })
    return questions, co_map

def get_tag_for_slot(slot_num, exam_type, part_c_choice, endsem_unit=None):
    if exam_type == "EndSem":
        mapping = {
            1:"1A", 2:"1A", 3:"2A", 4:"2A", 5:"3A", 6:"3A", 7:"4A", 8:"4A", 9:"5A", 10:"5A", 
            11:"1B", 12:"1B", 13:"2B", 14:"2B", 15:"3B", 16:"3B", 17:"4B", 18:"4B", 19:"5B", 20:"5B"
        }
        if slot_num in mapping: return mapping[slot_num]
        if slot_num in [21, 22]: return f"{endsem_unit}C"

    elif exam_type == "CAT 1":
        if slot_num in [1, 2]: return "1A"
        if slot_num in [3, 4]: return "2A"
        if slot_num == 5: return random.choice(["1A", "2A"])
        if slot_num in [6, 7]: return "1B"
        if slot_num in [8, 9]: return "2B"
        if slot_num in [10, 11]: return f"{part_c_choice}C"

    elif exam_type == "CAT 2":
        if slot_num in [1, 2]: return "3A"
        if slot_num in [3, 4]: return "4A"
        if slot_num == 5: return random.choice(["3A", "4A"])
        if slot_num in [6, 7]: return "3B"
        if slot_num in [8, 9]: return "4B"
        if slot_num in [10, 11]: return f"{part_c_choice}C"

    return None

def assemble_doc(template_bytes, selected_map, co_desc, mode):
    doc = Document(BytesIO(template_bytes))
    for table in doc.tables:
        for row in table.rows:
            txt = "".join([c.text for c in row.cells]).strip()
            for i in range(1, 6):
                if txt.startswith(f"CO{i}") and len(row.cells) >= 3:
                    row.cells[2].text = co_desc.get(str(i), f"Course Outcome {i}")

    for (ti, ri, ci), q in selected_map.items():
        table = doc.tables[ti]
        row = table.rows[ri]
        q_target, co_target = row.cells[1], row.cells[2]
        k_target = row.cells[3] if len(row.cells) > 3 else None
        
        if q == "TAG_NOT_FOUND":
            q_target.text = "Error: Tag missing in Bank"
        else:
            replace_cell_with_cell(q_target, q["q_cell"])
            if q.get("images"):
                for img_blob in q["images"]:
                    img_blob.seek(0)
                    p = q_target.add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_blob, width=Inches(2.0))
            if mode == "Faculty" and q.get("ans_text"):
                p = q_target.add_paragraph()
                p.add_run("\nAnswer: ").bold = True
                p.add_run(q["ans_text"])
            co_target.text = f"CO{q['tag'][0]}"
            if k_target: k_target.text = q.get('bloom', '')
            
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# -----------------------
# Streamlit App
# -----------------------
st.set_page_config(page_title="QP Generator", layout="wide")
st.title("📑 QP Dashboard")

if "s_zip" not in st.session_state:
    st.session_state.s_zip = None
    st.session_state.f_zip = None

with st.sidebar:
    st.header("Files")
    t_file = st.file_uploader("Template", type="docx")
    b_file = st.file_uploader("Bank", type="docx")
    st.header("Settings")
    # MAX 3 SETS REQUIREMENT
    n_sets = st.number_input("Sets", 1, 3, 1)
    exam = st.radio("Exam", ["CAT 1", "CAT 2", "EndSem"])
    pc, uc = None, None
    if exam == "EndSem": 
        uc = st.selectbox("Unit Part C", [1,2,3,4,5])
    elif exam == "CAT 1": 
        pc = st.selectbox("Unit Part C", [1,2])
    else: 
        pc = st.selectbox("Unit Part C", [3,4])

if st.button("🚀 Generate"):
    if t_file and b_file:
        with st.spinner("Processing..."):
            questions, co_desc = extract_bank_data(b_file)
            t_bytes = t_file.read()
            doc_t = Document(BytesIO(t_bytes))
            slots = []
            for ti, table in enumerate(doc_t.tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        m = NUM_PREFIX_RE.match(cell.text.strip())
                        if m: slots.append({"ti": ti, "ri": ri, "ci": ci, "sn": int(m.group(1))})
            
            s_buf, f_buf = BytesIO(), BytesIO()
            # GLOBAL TRACKING TO PREVENT REPETITION ACROSS SETS
            global_used = set()
            
            with zipfile.ZipFile(s_buf, "w") as sz, zipfile.ZipFile(f_buf, "w") as fz:
                for i in range(n_sets):
                    current_set_selection = {}
                    by_tag = defaultdict(list)
                    for q in questions: by_tag[q["tag"]].append(q)
                    
                    for s in slots:
                        tag = get_tag_for_slot(s["sn"], exam, pc, uc)
                        # Filter out questions already used in ANY previous set
                        pool = [q for q in by_tag.get(tag, []) if q["id"] not in global_used]
                        
                        if not pool:
                            # Fallback if the bank runs out of unique questions
                            chosen = "TAG_NOT_FOUND"
                        else:
                            chosen = random.choice(pool)
                            global_used.add(chosen["id"])
                            
                        current_set_selection[(s["ti"], s["ri"], s["ci"])] = chosen
                    
                    sz.writestr(f"Set_{i+1}_Student.docx", assemble_doc(t_bytes, current_set_selection, co_desc, "Student"))
                    fz.writestr(f"Set_{i+1}_Faculty.docx", assemble_doc(t_bytes, current_set_selection, co_desc, "Faculty"))
            
            st.session_state.s_zip = s_buf.getvalue()
            st.session_state.f_zip = f_buf.getvalue()
            st.success("Documents Ready! No questions are repeated across sets.")

if st.session_state.s_zip:
    c1, c2 = st.columns(2)
    c1.download_button("Download Student ZIP", st.session_state.s_zip, "Student.zip")
    c2.download_button("Download Faculty ZIP", st.session_state.f_zip, "Faculty.zip")
